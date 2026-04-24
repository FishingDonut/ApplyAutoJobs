import typer
import json
import os
import click
from rich.console import Console
from rich.table import Table
from .db.database import create_db_and_tables, get_session
from .db.models import Vaga, Perfil
from .db.migrate_csv import migrate as run_migration
from .scrapers.search import search_and_save
from .core.ai import GeminiEngine
from .bots.gupy import GupyBot
from .bots.linkedin import LinkedInBot
from sqlmodel import select
from docxtpl import DocxTemplate
from docx import Document
import time

app = typer.Typer(help="JobSpy AI - Assistente de Carreira no Terminal")
console = Console()

TEMPLATE_PATH = "templates/template-2.0.docx"
OUTPUT_DIR = "meus_curriculos"
PROFILE_PATH = "perfil.json"

from .core.match import MatchEngine

# --- LÓGICA COMPARTILHADA (CLI + TUI) ---

def generate_resume_logic(vaga_id: int, logger=None):
    """Lógica para gerar apenas o currículo personalizado e recalcular o match."""
    def log(msg):
        if logger: logger(msg)
        else: console.print(msg)

    with get_session() as session:
        vaga = session.get(Vaga, vaga_id)
        if not vaga:
            log(f"[bold red]Vaga ID {vaga_id} não encontrada![/bold red]")
            return False

        log(f"[yellow]Gerando currículo personalizado para: {vaga.titulo} na {vaga.empresa}...[/yellow]")
        
        # Extrair texto do template
        if not os.path.exists(TEMPLATE_PATH):
            log(f"[bold red]Template não encontrado em {TEMPLATE_PATH}[/bold red]")
            return False

        doc_base = Document(TEMPLATE_PATH)
        texto_base = "\n".join([p.text for p in doc_base.paragraphs if p.text])
        
        # IA Adaptação
        ai = GeminiEngine()
        res_raw = ai.adaptar_curriculo(vaga.descricao, texto_base)
        
        if not res_raw:
            log("[bold red]❌ A IA não retornou dados para adaptação.[/bold red]")
            return False

        try:
            # Tenta limpar tags de markdown se a IA ainda as retornar
            if "```json" in res_raw:
                res_raw = res_raw.split("```json")[1].split("```")[0].strip()
            elif "```" in res_raw:
                parts = res_raw.split("```")
                if len(parts) >= 3: res_raw = parts[1].strip()
            
            # Garante que temos apenas o objeto JSON caso haja texto extra
            if res_raw and not res_raw.startswith("{"):
                start = res_raw.find("{")
                end = res_raw.rfind("}")
                if start != -1 and end != -1: res_raw = res_raw[start:end+1]

            if not res_raw: raise ValueError("JSON vazio após processamento")

            dados = json.loads(res_raw)
            contexto = dados.get('adaptacao', {}) or dados
        except Exception as e:
            log(f"[bold red]❌ Erro no JSON da IA: {e}[/bold red]")
            log(f"[dim]Raw: {res_raw[:100]}...[/dim]")
            return False
        
        # Salvar Docx
        if not os.path.exists(OUTPUT_DIR): os.makedirs(OUTPUT_DIR)
        nome_arq = f"{vaga_id}_{vaga.empresa.replace(' ', '_')}.docx"
        caminho_final = os.path.abspath(os.path.join(OUTPUT_DIR, nome_arq))
        
        doc = DocxTemplate(TEMPLATE_PATH)
        doc.render(contexto)
        doc.save(caminho_final)
        
        # Recálculo Matemático do Match
        doc_final = Document(caminho_final)
        texto_final = "\n".join([p.text for p in doc_final.paragraphs if p.text])
        
        engine = MatchEngine()
        novo_score = engine.recalcular_match_com_curriculo(texto_final, vaga.descricao)
        
        log(f"[green]✅ Currículo salvo em: {caminho_final}[/green]")
        log(f"[bold cyan]📊 Match Recalculado (Matemático): {vaga.match_score}% -> {novo_score}%[/bold cyan]")
        
        vaga.arquivo_curriculo = caminho_final
        vaga.match_score = novo_score
        vaga.status = "Currículo Gerado"
        session.add(vaga)
        session.commit()
        return True

def apply_logic(vaga_id: int, logger=None):
    """Lógica centralizada de aplicação para ser usada pela CLI e pela TUI."""
    def log(msg):
        if logger: logger(msg)
        else: console.print(msg)

    with get_session() as session:
        vaga = session.get(Vaga, vaga_id)
        if not vaga:
            log(f"[bold red]Vaga ID {vaga_id} não encontrada![/bold red]")
            return False

        # Se não tiver currículo, gera um agora
        if not vaga.arquivo_curriculo:
            if not generate_resume_logic(vaga_id, logger):
                return False
            # Recarregar vaga após commit no generate_resume_logic
            session.refresh(vaga)

        caminho_final = vaga.arquivo_curriculo

        # Automação
        sucesso = False
        if vaga.site == "Gupy":
            log("[yellow]Iniciando GupyBot...[/yellow]")
            bot = GupyBot()
            sucesso = bot.aplicar(vaga.link, caminho_final)
        elif vaga.site == "LinkedIn":
            log("[yellow]Iniciando LinkedInBot...[/yellow]")
            bot = LinkedInBot()
            sucesso = bot.aplicar(vaga.link, caminho_final)
        
        if sucesso:
            vaga.status = "Candidatado"
            session.add(vaga)
            session.commit()
            log("[bold green]🚀 Candidatura finalizada com sucesso![/bold green]")
            return True
        
        return False

# --- COMANDOS CLI ---

@app.command()
def login(plataforma: str = typer.Argument(..., help="Plataforma para login (gupy ou linkedin)")):
    """Abre o navegador para realizar o login manual e salvar a sessão."""
    from playwright.sync_api import sync_playwright
    user_data = ".sessao_gupy" if plataforma.lower() == "gupy" else ".sessao_linkedin"
    url = "https://portal.gupy.io/" if plataforma.lower() == "gupy" else "https://www.linkedin.com/"
    
    console.print(f"[bold yellow]Iniciando login na {plataforma.upper()}...[/bold yellow]")
    with sync_playwright() as p:
        context = p.chromium.launch_persistent_context(user_data_dir=user_data, headless=False, channel="chrome")
        page = context.new_page()
        page.goto(url)
        while len(context.pages) > 0: time.sleep(1)
    console.print(f"[bold green]Sessão salva![/bold green]")

@app.command()
def setup():
    """Inicializa o banco de dados MySQL."""
    create_db_and_tables()
    console.print("[green]Banco de dados pronto![/green]")

@app.command()
def search(termo: str, local: str = "remoto", limite: int = 20):
    """Busca novas vagas e as salva no banco de dados."""
    setup()
    novas = search_and_save(termo, local, True, limite)
    console.print(f"[bold green]Sucesso![/bold green] {novas} novas vagas encontradas.")

@app.command()
def list(limit: int = 10):
    """Lista as últimas vagas com detalhes do currículo."""
    with get_session() as session:
        vagas = session.exec(select(Vaga).order_by(Vaga.data_descoberta.desc()).limit(limit)).all()
        table = Table(title="Últimas Vagas")
        table.add_column("ID", style="cyan")
        table.add_column("Título", style="magenta")
        table.add_column("Empresa", style="green")
        table.add_column("Match", style="bold yellow")
        table.add_column("Currículo", style="dim")
        table.add_column("Status", style="white")
        
        for v in vagas:
            curr = "✅" if v.arquivo_curriculo else "❌"
            table.add_row(str(v.id), v.titulo, v.empresa, f"{v.match_score}%", curr, v.status)
        console.print(table)

@app.command()
def dashboard():
    """Abre o painel interativo (TUI)."""
    from .tui.dashboard import JobSpyDashboard
    JobSpyDashboard().run()

@app.command()
def generate(vaga_id: int):
    """Apenas gera o currículo personalizado e recalcular o match."""
    generate_resume_logic(vaga_id)

@app.command()
def apply(vaga_id: int):
    """Adapta o currículo (se necessário) e aplica para uma vaga."""
    apply_logic(vaga_id)

@app.command()
def profile():
    """Editor interativo de perfil para gerenciar suas habilidades e experiências."""
    if not os.path.exists(PROFILE_PATH):
        console.print("[yellow]Perfil não encontrado. Criando um novo...[/yellow]")
        perfil = {
            "dados_pessoais": {"nome_completo": "", "email": "", "telefone": "", "cidade": "", "estado": ""},
            "preferencias": {"pretensao_salarial_clt": 0, "modelo_trabalho": ["Remoto"]},
            "skills": "Python, SQL, Engenharia de Dados",
            "resumo": "Desenvolvedor experiente em busca de novos desafios.",
            "prompt_personalizado_respostas": "Seja direto e técnico."
        }
    else:
        with open(PROFILE_PATH, "r", encoding="utf-8") as f:
            perfil = json.load(f)

    # Garantir campos mínimos para o MatchEngine funcionar
    if "skills" not in perfil: perfil["skills"] = ""
    if "resumo" not in perfil: perfil["resumo"] = ""

    while True:
        console.rule("[bold blue]JobSpy Profile Editor[/bold blue]")
        console.print(f"[bold cyan]Nome:[/bold cyan] {perfil.get('dados_pessoais', {}).get('nome_completo', '---')}")
        console.print(f"[bold cyan]Skills:[/bold cyan] {perfil.get('skills', '---')}")
        console.print("-" * 30)
        console.print("1. [bold]Dados Pessoais[/bold] (Nome, Email, Tel, Localização)")
        console.print("2. [bold]Preferências[/bold] (Salário, Modelo de Trabalho)")
        console.print("3. [bold]Habilidades (Skills)[/bold] [dim]- Usado para Match %[/dim]")
        console.print("4. [bold]Resumo Profissional[/bold] [dim]- Usado para Match %[/dim]")
        console.print("5. [bold]Prompt da IA[/bold]")
        console.print("6. [bold]Abrir no Editor de Texto[/bold] (JSON Completo)")
        console.print("0. [bold green]Salvar e Sair[/bold green]")
        
        opcao = typer.prompt("\nEscolha uma opção", default="0")

        if opcao == "1":
            dp = perfil.get("dados_pessoais", {})
            dp["nome_completo"] = typer.prompt("Nome Completo", default=dp.get("nome_completo", ""))
            dp["email"] = typer.prompt("Email", default=dp.get("email", ""))
            dp["telefone"] = typer.prompt("Telefone", default=dp.get("telefone", ""))
            dp["cidade"] = typer.prompt("Cidade", default=dp.get("cidade", ""))
            dp["estado"] = typer.prompt("Estado", default=dp.get("estado", ""))
            perfil["dados_pessoais"] = dp
        elif opcao == "2":
            pref = perfil.get("preferencias", {})
            pref["pretensao_salarial_clt"] = float(typer.prompt("Pretensão Salarial CLT (R$)", default=str(pref.get("pretensao_salarial_clt", 0))))
            perfil["preferencias"] = pref
        elif opcao == "3":
            console.print("[dim]Dica: Liste suas tecnologias separadas por vírgula para melhor precisão no Match.[/dim]")
            perfil["skills"] = typer.prompt("Habilidades", default=perfil.get("skills", ""))
        elif opcao == "4":
            perfil["resumo"] = typer.prompt("Resumo Profissional", default=perfil.get("resumo", ""))
        elif opcao == "5":
            perfil["prompt_personalizado_respostas"] = typer.prompt("Prompt para a IA (como ela deve agir)", default=perfil.get("prompt_personalizado_respostas", ""))
        elif opcao == "6":
            console.print("[yellow]Abrindo arquivo JSON... Salve e feche para retornar.[/yellow]")
            click.edit(filename=PROFILE_PATH)
            # Recarregar após edição manual
            try:
                with open(PROFILE_PATH, "r", encoding="utf-8") as f:
                    perfil = json.load(f)
            except Exception as e:
                console.print(f"[bold red]Erro ao ler o arquivo após edição: {e}[/bold red]")
        elif opcao == "0":
            break
    
    with open(PROFILE_PATH, "w", encoding="utf-8") as f:
        json.dump(perfil, f, indent=2, ensure_ascii=False)
    
    console.print("[bold green]✅ Perfil atualizado e salvo com sucesso![/bold green]")

if __name__ == "__main__":
    app()
